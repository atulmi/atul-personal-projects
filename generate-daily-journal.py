# Purpose of this program: generate a docx for each month of the year, containing a page (with 15 blank bullet points) for each day of that month
# Example usage: python3 generate-daily-journal.py

# Result: 12 docx files

from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Cm
from docx.shared import RGBColor

year = datetime.now().year

start_date = datetime(year, 1, 1)
end_date = datetime(year, 12, 31)
date_list = [start_date + timedelta(days=i) for i in range((end_date - start_date).days + 1)] # list of all dates in the current year

# Initialize the current date to the start date
current_date = start_date
current_month = ''
document = Document()

# Loop through the list of all dates for this year
for date in date_list:
    # Get the month name and year
    month_name = date.strftime("%B")  # Full month name
    year = date.year

    # See if we've looped to a new month... if so, update the current_month
    if current_month != month_name:
        document = Document()
        current_month = month_name

    
    formatted_date = date.strftime("%b %d, %Y")  # format as "Jan 01, 2025 or Jan 30, 2025, etc"    
    paragraph = document.add_paragraph()

    temp_runner = paragraph.add_run("" + formatted_date)
    temp_runner.bold = True
    temp_runner.font.color.rgb = RGBColor(255, 0, 0)  # Make date header red

    for i in range(15):
        newParagraph = document.add_paragraph(style='ListBullet')
        newParagraph.add_run("")
        document.add_paragraph() # Empty line under bullet point to write stuff down

    document.add_page_break()


    sections = document.sections

    # set margins of docx file to very low margins (we want to maximize portion of page used for text)
    for section in sections:
       section.top_margin = Cm(0.1)
       section.bottom_margin = Cm(0.1)
       section.left_margin = Cm(0.4)
       section.right_margin = Cm(0.1)

    # save to file with name corresponding to current_month (April, September, etc)
    document.save(current_month + ".docx")
