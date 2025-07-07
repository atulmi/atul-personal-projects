# Goal of this program: generate a docx note-taking template for a set of PDF books in the given directory (passed as 1st argument on command-line)
# Each docx will be populated with entries/headings for each block of "n" pages until final page count of book is reached (number of pages per block is provided as 2nd argument)

# Example usage: "python3 generate-book-notes-page-blocks.py Example-PDF-Directory/ 20"

# Ex: if n=20, the docx will have page block entries for pages "1-20", "21-40", "41-60"... "560-567", etc, all on separate lines. The page block entry/heading will be bolded
# Ex: if n=1, the docx will have headings like "Page 1, Page 2, Page 3", etc, all on separate lines. The page entry/heading will be bolded
# The docx templates will be created in CURRENT directory (not the given one), so the given directory (with PDF books) doesn't get cluttered with the note-taking templates

# The program is fully local, so the PDF files won't be uploaded to the cloud unless we explicitly upload them (or have some sort of automated backup service)
# Not all PDFs have a clear concept of chapters (they might have "sections" or "parts"), so I opted to create entries for each set of "n" pages, rather than 1 entry for each chapter
# This program only supports PDF files. If there are books in EPUB or other book formats, use Calibre or some other program to convert to PDF before running this program

from docx import Document
from docx.shared import Inches, Cm
import sys
import os
import fitz
from pathlib import Path


document = Document()

if len(sys.argv) < 3:
    print("2 arguments are needed for this program: the directory containing PDF files, and the page block count")
    sys.exit()

source_dir = sys.argv[1] # given directory with PDF files to open
numPagesPerBlock = sys.argv[2] # number of pages per block

if str.isdigit(numPagesPerBlock) == False:
    print("The provided 2nd argument (page block count) is not a number")
    sys.exit()

pdf_files = []
pdf_filenames = []
pdf_names_and_pages = []


# navigate recursively through all PDFs from given directory, create array of files
for path in Path(source_dir).rglob('*.pdf'):
    pdf_files.append(path.resolve())
    pdf_filenames.append(path.name)


# Loop through array of PDF files and get page counts
for index, pdf in enumerate(pdf_files):
    file = fitz.open(pdf)

    totalPages = len(file)

    toWrite1 = str(pdf_filenames[index])
    toWrite2 = totalPages

    # add to array of tuples (each tuple has PDF file name & its page count)
    pdf_names_and_pages.append((toWrite1, toWrite2))

    file.close()



# Return page blocks (e.g. "1-20", "21-40", "560-567"), up until final page count
def get_ranges_up_to_number(n):
    ranges = []
    for i in range(1, n + 1, int(float(numPagesPerBlock))):
        ranges.append((i, min(i + int(float(numPagesPerBlock)) - 1, n)))
    return ranges



# Loop through array of tuples (each tuple with book PDF name and its page count)
for i in range(len(pdf_names_and_pages)):
    document = Document()
    line = str(pdf_names_and_pages[i][0]) + "\t" + str(pdf_names_and_pages[i][1])

    title = line.split("\t")[0].replace(".pdf", "").replace("/", "")
    pages = line.split("\t")[1]

    # If no errors in opening PDF (page count shouldn't be 'nan') and a note doesn't already exist, create a note for this PDF title
    if pages != 'nan' and os.path.exists(title + ".docx") == False:

        pageCount = int(float(pages))

        # if page block size is over 1, return headings like "1-20", "1-5", etc, depending on the page block size
        if int(float(numPagesPerBlock)) > 1:
            pageBlocksOfN = get_ranges_up_to_number(pageCount)

            # add page blocks (e.g. "1-20", "21-40") into docx
            for j in pageBlocksOfN:
                paragraph = document.add_paragraph()

                # Create "page block" heading/entry
                temp_runner = paragraph.add_run("Pages " + str(j[0]) + "-" + str(j[1]))
                temp_runner.bold = True

                # Add colon after "page block" heading, everything after heading is unbolded
                temp_runner_new = paragraph.add_run(": ")
                temp_runner_new.bold = False

                document.save(title + ".docx")

        # if page block size is just 1, make heading "Page 1, Page 2..." rather than a range like "Pages 1-2"
        else:
            for j in range(1, pageCount + 1):
                paragraph = document.add_paragraph()

                temp_runner = paragraph.add_run("Page " + str(j))
                temp_runner.bold = True

                # Add colon after "page block" heading, everything after heading is unbolded
                temp_runner_new = paragraph.add_run(": ")
                temp_runner_new.bold = False

                document.save(title + ".docx")



        sections = document.sections

        # set margins of docx file to very low margins (we want to maximize portion of page used for text)
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(0.5)
            section.right_margin = Cm(0.5)


        document.save(title + ".docx")
