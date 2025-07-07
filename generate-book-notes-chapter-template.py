# Goal of this program: generate a docx note-taking template for a set of books
# For each book, a docx will be generated with "Intro, Chapter 1, Chapter 2... Conclusion" all on separate lines
# As a command-line argument, we pass a spreadsheet, where each row has 2 items: book title & its number of chapters
# The program is fully local, so the generated docx files won't be uploaded to the cloud unless we explicitly upload them (or have some sort of automated backup service)
# Usage: "python3 generate-book-notes-chapter-template.py test-spreadsheet.xlsx"
# Result: 1 note-taking template for each book listed in above spreadsheet will be created in CURRENT directory, unless a docx with that book already exists (we don't want to overwrite existing book note documents!)

from docx import Document
import pandas as pd
import sys
import os


document = Document()

df = pd.read_excel(sys.argv[1])

# Loop through all book titles listed in spreadsheet
for i in range(len(df)):
    document = Document()
    line = str(df.iloc[i, 0]) + "\t" + str(df.iloc[i, 1])

    title = line.split("\t")[0]
    chapters = line.split("\t")[1]

    # Start generating note template docx, if "chapters" cell is a number, and if note for current title doesn't exist (in current directory)
    if chapters != 'nan' and os.path.exists(title + ".docx") == False:

        paragraph = document.add_paragraph()
        intro_runner = paragraph.add_run("Introduction")
        intro_runner.bold = True

        # Create some empty lines after "Introduction" to type in text
        intro_runner_new = paragraph.add_run(": \n\n\n")
        intro_runner_new.bold = False
        document.save(title + ".docx")

        # Loop through number of chapters, create an entry in docx for each
        for j in range(0, int(float(chapters))):
              paragraph = document.add_paragraph()
              temp_runner = paragraph.add_run("Chapter " + str(j + 1))
              temp_runner.bold = True

              # Create some empty lines after chapter heading, to type in text
              temp_runner_new = paragraph.add_run(": \n\n\n")
              temp_runner_new.bold = False

              document.save(title + ".docx")


        paragraph = document.add_paragraph()
        conclusion_runner = paragraph.add_run("Conclusion")
        conclusion_runner.bold = True
        
        conclusion_runner_new = paragraph.add_run(": \n\n\n")
        conclusion_runner_new.bold = False

        document.save(title + ".docx")

