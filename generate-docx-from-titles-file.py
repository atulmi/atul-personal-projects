from docx import Document
from docx.shared import Inches, Cm, RGBColor
from os import sys

if(len(sys.argv) > 1):
    filename = sys.argv[1]

else:
    raise ValueError('Please provide a file as an argument')

with open(filename, 'r') as file:
    lines = file.readlines()
    
    # Number of titles per document
    titles_per_doc = 100
    file_count = 1  # Start file count

    for i in range(0, len(lines), titles_per_doc):
        document = Document()  # Create a new Document
        for line in lines[i:i + titles_per_doc]:
            title = line.strip()  # Clean up the line
            if title:  # Check if the line is not empty
                paragraph = document.add_paragraph()

                temp_runner = paragraph.add_run(title)
                temp_runner.bold = True
                temp_runner.font.color.rgb = RGBColor(255, 0, 0)  # Make header red

                document.add_page_break()  # Add a page break
                
                
                #paragraph2 = document.add_paragraph()

                #temp_runner1 = paragraph2.add_run(title)
                #temp_runner1.bold = True
                #temp_runner1.font.color.rgb = RGBColor(255, 0, 0)  # Make header red

                #document.add_page_break()  # Add a page break


                sections = document.sections

                # set margins of docx file to very low margins (we want to maximize portion of page used for text)
                for section in sections:
                    section.top_margin = Cm(0.1)
                    section.bottom_margin = Cm(0.1)
                    section.left_margin = Cm(0.1)
                    section.right_margin = Cm(0.1)



        # Save the document
        document.save(filename.split('.')[0] + '-' + str(file_count) + ".docx")
        file_count += 1  # Increment file count for next document
